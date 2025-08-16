import os
import uuid
import logging
from typing import List, Tuple
from pathlib import Path

import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

from config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
        )
    
    def extract_text(self, file_path: str, filename: str) -> str:
        """Extract text from various file formats."""
        file_ext = Path(filename).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_ext == '.docx':
                return self._extract_docx_text(file_path)
            elif file_ext == '.txt':
                return self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def chunk_text(self, text: str, metadata: dict = None) -> List[LangchainDocument]:
        """Split text into chunks for vector storage."""
        if metadata is None:
            metadata = {}
        
        # Create documents with metadata
        doc = LangchainDocument(page_content=text, metadata=metadata)
        chunks = self.text_splitter.split_documents([doc])
        
        # Add chunk index to metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                "chunk_index": i,
                "total_chunks": len(chunks)
            })
        
        return chunks